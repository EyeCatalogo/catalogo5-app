# streamlit_app.py
import streamlit as st
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import pandas as pd
import tempfile
import requests
from io import BytesIO
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches

# -------------------------
# Config y requirements
# -------------------------
st.set_page_config(page_title="Catálogo - Google Sheets → PDF/DOCX", layout="wide")
st.title("🛍️ Generador de Catálogo (Google Sheets)")

# Requisitos sugeridos (requirements.txt):
# streamlit
# gspread
# oauth2client
# pandas
# requests
# reportlab
# python-docx
# openpyxl

# -------------------------
# Helpers: conectar a Google
# -------------------------
def guardar_json_temp(uploaded_json):
    """Guarda el JSON subido en un archivo temporal y devuelve la ruta."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".json") as temp:
        temp.write(uploaded_json.read())
        return temp.name

def conectar_gspread(json_path):
    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive"
    ]
    creds = ServiceAccountCredentials.from_json_keyfile_name(json_path, scope)
    client = gspread.authorize(creds)
    return client

# -------------------------
# Crear template en Google Sheets
# -------------------------
def crear_template_en_sheets(client, spreadsheet_name="Catalogo"):
    try:
        try:
            spreadsheet = client.open(spreadsheet_name)
            created = False
        except gspread.SpreadsheetNotFound:
            spreadsheet = client.create(spreadsheet_name)
            created = True
        # crear/obtener worksheet "Catalogo"
        try:
            worksheet = spreadsheet.worksheet("Catalogo")
        except gspread.WorksheetNotFound:
            worksheet = spreadsheet.add_worksheet(title="Catalogo", rows="200", cols="10")

        headers = ["categoria", "nombre", "descripcion", "precio", "stock", "imagen"]
        worksheet.update("A1", [headers])

        datos_demo = [
            ["Electrónica", 'Televisor Samsung 40"', "Smart TV 40 pulgadas", "250", "8",
             "https://drive.google.com/file/d/10VB9sF9j6FXvRRCFM4t7t7idBkz9KARc/view?usp=sharing"],
            ["Electrónica", "Laptop HP 15\"", "15'' RAM 8GB", "500", "4",
             "https://drive.google.com/file/d/1bVgLB1ps02AYEzoPEQnsxB5lUfO9dJDY/view?usp=sharing"],
            ["Hogar", "Silla ergonómica", "Con soporte lumbar", "80", "12",
             "https://drive.google.com/file/d/1vV9AD4S1zowIrW-rtTq-6Zb8xRj1zPqg/view?usp=sharing"],
            ["Ropa", "Camiseta Polo", "Algodón premium", "30", "30",
             "https://drive.google.com/file/d/1-7LrG5cwqQ1bQhU3F2_t5GCKuVWkQUtw/view?usp=sharing"]
        ]
        worksheet.update("A2", datos_demo)

        return spreadsheet
    except Exception as e:
        st.error(f"Error creando template: {e}")
        return None

# -------------------------
# Cargar datos desde Google Sheets
# -------------------------
def cargar_datos_google(json_path, spreadsheet_name="Catalogo"):
    try:
        client = conectar_gspread(json_path)
        try:
            spreadsheet = client.open(spreadsheet_name)
        except gspread.SpreadsheetNotFound:
            st.error(f"No existe el Google Sheet llamado '{spreadsheet_name}'. Puedes crear el template desde la app.")
            return None, client
        try:
            worksheet = spreadsheet.worksheet("Catalogo")
        except gspread.WorksheetNotFound:
            st.error("La hoja 'Catalogo' no existe. Crea el template en el menú lateral.")
            return None, client
        data = worksheet.get_all_records()
        df = pd.DataFrame(data)
        return df, client
    except Exception as e:
        st.error(f"Error al conectar/leer Google Sheets: {e}")
        return None, None

# -------------------------
# Util: descargar imagen desde URL (devuelve BytesIO o None)
# -------------------------
def descargar_imagen_bytes(url):
    try:
        if not url:
            return None
        url = str(url).strip()
        if url == "" or url.lower() == "nan":
            return None
        # Manejar Google Drive share links
        if "drive.google.com" in url:
            if "/d/" in url:
                file_id = url.split("/d/")[1].split("/")[0]
                url = f"https://drive.google.com/uc?export=view&id={file_id}"
            elif "id=" in url:
                file_id = url.split("id=")[1].split("&")[0]
                url = f"https://drive.google.com/uc?export=view&id={file_id}"
        resp = requests.get(url, timeout=10)
        if resp.status_code == 200 and "image" in resp.headers.get("content-type", ""):
            return BytesIO(resp.content)
    except Exception:
        return None
    return None

# -------------------------
# 1) Catálogo Final (PDF parametrizable)
# -------------------------
def generar_catalogo_pdf(dataframe, tema_color_hex="#2E86C1", mini_logo_bytes=None, portada_info=None):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()
    tema_color = colors.HexColor(tema_color_hex)

    # Portada opcional (si portada_info dict con title/subtitle/logo_bytes)
    if portada_info:
        title = portada_info.get("title", "")
        subtitle = portada_info.get("subtitle", "")
        logo_b = portada_info.get("logo_bytes", None)
        # Logo
        story.append(Spacer(1, 2 * cm))
        if logo_b:
            try:
                story.append(RLImage(logo_b, width=6*cm, height=6*cm))
                story.append(Spacer(1, 0.5*cm))
            except Exception:
                pass
        if title:
            story.append(Paragraph(f"<b>{title}</b>", ParagraphStyle(name="PortTitle", fontSize=22, alignment=1, textColor=tema_color)))
            story.append(Spacer(1, 0.2*cm))
        if subtitle:
            story.append(Paragraph(subtitle, ParagraphStyle(name="PortSub", fontSize=12, alignment=1)))
            story.append(Spacer(1, 0.5*cm))
        fecha = datetime.now().strftime("%d %B %Y")
        story.append(Paragraph(f"<i>Generado: {fecha}</i>", ParagraphStyle(name="PortDate", fontSize=9, alignment=1, textColor=colors.grey)))
        story.append(PageBreak())

    styles.add(ParagraphStyle(name="CategoriaTitle", fontSize=16, leading=18, spaceAfter=8, textColor=tema_color))
    styles.add(ParagraphStyle(name="ProductoTitle", fontSize=12, leading=14, alignment=1, textColor=colors.HexColor("#212F3D")))
    styles.add(ParagraphStyle(name="ProductoText", fontSize=10, leading=12))

    if "categoria" not in dataframe.columns:
        st.warning("La columna 'categoria' no existe. Se generará un listado sin agrupar.")
        categorias = ["Todos"]
        grouped = [("Todos", dataframe)]
    else:
        grouped = list(dataframe.groupby("categoria"))

    productos_por_fila = 2

    for categoria, grupo in grouped:
        # Título de categoría
        story.append(Paragraph(f"{categoria}", styles["CategoriaTitle"]))
        story.append(Spacer(1, 0.2*cm))

        fila = []
        celdas = []
        for _, row in grupo.iterrows():
            nombre = str(row.get("nombre", row.get("Nombre", "")))
            descripcion = str(row.get("descripcion", row.get("Descripcion", ""))) if "descripcion" in row else ""
            precio = str(row.get("precio", row.get("Precio", "")))
            stock = str(row.get("stock", row.get("Stock", "")))
            imagen_url = row.get("imagen", "")

            img_bytes = descargar_imagen_bytes(imagen_url)
            if img_bytes:
                img_flow = RLImage(img_bytes, width=5*cm, height=5*cm)
            else:
                img_flow = Paragraph("🖼️ Imagen no disponible", styles["ProductoText"])

            elementos = [
                img_flow,
                Paragraph(f"<b>{nombre}</b>", styles["ProductoTitle"]),
                Paragraph(descripcion, styles["ProductoText"]),
                Paragraph(f"Precio: ${precio}", styles["ProductoText"]),
                Paragraph(f"Stock: {stock}", styles["ProductoText"]),
            ]
            if mini_logo_bytes:
                try:
                    mini_img = RLImage(mini_logo_bytes, width=0.8*cm, height=0.8*cm)
                    elementos.append(mini_img)
                except Exception:
                    pass

            ficha = Table([[e] for e in elementos], colWidths=[6.8*cm])
            ficha.setStyle(TableStyle([
                ("ALIGN", (0,0), (-1,-1), "CENTER"),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("BOX", (0,0), (-1,-1), 0.25, colors.grey),
                ("TOPPADDING", (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 6)
            ]))

            fila.append(ficha)
            if len(fila) == productos_por_fila:
                celdas.append(fila)
                fila = []

        if fila:
            celdas.append(fila)

        if celdas:
            tabla = Table(celdas, colWidths=[9*cm]*productos_por_fila)
            tabla.setStyle(TableStyle([
                ("ALIGN", (0,0), (-1,-1), "CENTER"),
                ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
                ("TOPPADDING", (0,0), (-1,-1), 10),
                ("BOTTOMPADDING", (0,0), (-1,-1), 10),
            ]))
            story.append(tabla)

        story.append(PageBreak())

    doc.build(story)
    buffer.seek(0)
    return buffer

# -------------------------
# 2) Mockup Visual (zonas de edición)
# -------------------------
def generar_mockup_visual():
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    story = []
    styles = getSampleStyleSheet()

    story.append(Paragraph("📘 Guía Visual - Mockup de Catálogo", styles["Title"]))
    story.append(Spacer(1, 0.5*cm))

    zonas = [
        ("Zona: Logo / Cabecera", colors.Color(0.9,0.95,1)),
        ("Zona: Título de categoría", colors.Color(0.95,1,0.95)),
        ("Zona: Ficha de producto (imagen + datos)", colors.Color(0.98,0.98,0.98)),
        ("Zona: Mini logo (opcional)", colors.Color(1,0.98,0.9)),
    ]

    for label, bgcolor in zonas:
        t = Table([[Paragraph(label, styles["Normal"])]], colWidths=[16*cm], rowHeights=[2.2*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), bgcolor),
            ("BOX", (0,0), (-1,-1), 1, colors.grey),
            ("ALIGN", (0,0), (-1,-1), "CENTER"),
            ("VALIGN", (0,0), (-1,-1), "MIDDLE")
        ]))
        story.append(t)
        story.append(Spacer(1, 0.4*cm))

    doc.build(story)
    buffer.seek(0)
    return buffer

# -------------------------
# 3) Versión editable DOCX
# -------------------------
def generar_version_editable_docx(dataframe):
    doc = Document()
    doc.add_heading("Catálogo de Productos", level=1)
    doc.add_paragraph(f"Generado: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph("")

    for _, row in dataframe.iterrows():
        nombre = str(row.get("nombre", row.get("Nombre", "")))
        categoria = str(row.get("categoria", row.get("Categoria", "")))
        descripcion = str(row.get("descripcion", row.get("Descripcion", ""))) if "descripcion" in row else ""
        precio = str(row.get("precio", row.get("Precio", "")))
        stock = str(row.get("stock", row.get("Stock", "")))
        imagen_url = row.get("imagen", "")

        doc.add_heading(nombre, level=2)
        doc.add_paragraph(f"Categoría: {categoria}")
        if descripcion:
            doc.add_paragraph(descripcion)
        doc.add_paragraph(f"Precio: ${precio}")
        doc.add_paragraph(f"Stock: {stock}")

        # intentar añadir imagen (si existe)
        img_bytes = descargar_imagen_bytes(imagen_url)
        if img_bytes:
            try:
                doc.add_picture(img_bytes, width=Inches(2.5))
            except Exception:
                pass
        doc.add_paragraph("")  # separador

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# -------------------------
# Interfaz Streamlit
# -------------------------
st.sidebar.header("Conectar Google Sheets / Template")
uploaded_json = st.sidebar.file_uploader("Sube credenciales (JSON)", type=["json"])
spreadsheet_name = st.sidebar.text_input("Nombre del Google Sheet", value="Catalogo")

if uploaded_json:
    json_path = guardar_json_temp(uploaded_json)
    client = None
    try:
        client = conectar_gspread(json_path)
        st.sidebar.success("✅ Conexión preparada")
    except Exception as e:
        st.sidebar.error(f"Error con credenciales: {e}")
        client = None

    if client:
        if st.sidebar.button("Crear template en Google Sheets"):
            crear_template_en_sheets(client, spreadsheet_name)
        if st.sidebar.button("Cargar datos (hoja 'Catalogo')"):
            df, _ = cargar_datos_google(json_path, spreadsheet_name)
            if df is not None and not df.empty:
                st.success("✅ Datos cargados")
                st.dataframe(df)
                st.session_state["df"] = df
            else:
                st.warning("La hoja 'Catalogo' está vacía o no tiene el formato esperado.")

# -------------------------
# Botones principales (3 outputs)
# -------------------------
if "df" in st.session_state:
    df = st.session_state["df"]
    st.markdown("### 📄 Generar archivos")
    # inputs para personalización
    tema_color = st.color_picker("Color principal del PDF (tema)", "#2E86C1")
    portada_title = st.text_input("Título de portada (opcional)", value="Catálogo de Productos")
    portada_sub = st.text_input("Subtítulo de portada (opcional)", value="Lista de productos")
    logo_file = st.file_uploader("Sube logo de portada (opcional)", type=["png","jpg"])
    mini_logo_file = st.file_uploader("Sube mini-logo (opcional)", type=["png","jpg"])

    # preparar bytes
    logo_bytes = None
    if logo_file:
        logo_bytes = BytesIO(logo_file.read())
    mini_logo_bytes = None
    if mini_logo_file:
        mini_logo_bytes = BytesIO(mini_logo_file.read())

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("📘 Catálogo Final (PDF)"):
            portada_info = {"title": portada_title, "subtitle": portada_sub, "logo_bytes": logo_bytes}
            pdf = generar_catalogo_pdf(df, tema_color_hex=tema_color, mini_logo_bytes=mini_logo_bytes, portada_info=portada_info)
            st.success("Catálogo PDF generado")
            st.download_button("⬇️ Descargar PDF Final", data=pdf, file_name="catalogo_final.pdf", mime="application/pdf")

    with col2:
        if st.button("🧱 Mockup Visual (PDF)"):
            pdfm = generar_mockup_visual()
            st.success("Mockup PDF generado")
            st.download_button("⬇️ Descargar Mockup", data=pdfm, file_name="mockup_visual.pdf", mime="application/pdf")

    with col3:
        if st.button("✏️ Versión Editable (DOCX)"):
            docx_b = generar_version_editable_docx(df)
            st.success("DOCX generado")
            st.download_button("⬇️ Descargar DOCX editable", data=docx_b, file_name="catalogo_editable.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

else:
    st.info("Sube las credenciales y carga la hoja 'Catalogo' para generar archivos.")
